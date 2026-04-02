"""Report generation service — PDF, DOCX, and JSON reports."""
import io
import json
import logging
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class ReportService:
    """Generate downloadable reports in PDF, DOCX, and JSON formats."""

    def generate_report(self, analysis: dict, bugs: list, insights: list,
                        report_type: str = "full", format: str = "pdf") -> bytes:
        """Generate a report in the specified format."""
        if format == "pdf":
            return self._generate_pdf(analysis, bugs, insights, report_type)
        elif format == "docx":
            return self._generate_docx(analysis, bugs, insights, report_type)
        elif format == "json":
            return self._generate_json(analysis, bugs, insights, report_type)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_pdf(self, analysis: dict, bugs: list, insights: list, report_type: str) -> bytes:
        """Generate PDF report using ReportLab."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'],
                                      fontSize=24, textColor=colors.HexColor('#1a1a2e'))
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'],
                                        fontSize=16, textColor=colors.HexColor('#16213e'),
                                        spaceAfter=12)
        body_style = ParagraphStyle('CustomBody', parent=styles['Normal'],
                                     fontSize=10, spaceAfter=6)

        elements = []

        # Title Page
        elements.append(Spacer(1, 2*inch))
        elements.append(Paragraph("🔷 DevInsight", title_style))
        elements.append(Paragraph("Code Quality Intelligence Report", heading_style))
        elements.append(Spacer(1, 0.25*inch))
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style))
        elements.append(Paragraph(f"Report Type: {report_type.replace('_', ' ').title()}", body_style))
        elements.append(PageBreak())

        if report_type in ("full", "insights_only"):
            # Executive Summary
            elements.append(Paragraph("Executive Summary", heading_style))
            summary_data = [
                ["Metric", "Value"],
                ["Total Files", str(analysis.get("total_files", 0))],
                ["Total Lines of Code", str(analysis.get("total_lines", 0))],
                ["Average Complexity", f"{analysis.get('avg_complexity', 0):.1f}"],
                ["Maximum Complexity", f"{analysis.get('max_complexity', 0):.1f}"],
                ["Technical Debt", f"{analysis.get('technical_debt_hours', 0):.1f} hours"],
                ["Risk Level", analysis.get("risk_level", "low").upper()],
                ["Risk Score", f"{analysis.get('risk_score', 0):.1f}/100"],
            ]

            table = Table(summary_data, colWidths=[3*inch, 3*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a1a2e')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0f0f5')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            elements.append(table)
            elements.append(Spacer(1, 0.5*inch))

            # Language breakdown
            lang_data = analysis.get("language_breakdown", {})
            if lang_data:
                elements.append(Paragraph("Language Breakdown", heading_style))
                lang_table_data = [["Language", "Files", "Lines"]]
                for lang, data in lang_data.items():
                    lang_table_data.append([lang.capitalize(), str(data["files"]), str(data["lines"])])

                lang_table = Table(lang_table_data, colWidths=[2*inch, 2*inch, 2*inch])
                lang_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#16213e')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                elements.append(lang_table)
                elements.append(Spacer(1, 0.5*inch))

        if report_type in ("full", "insights_only"):
            # AI Insights
            elements.append(Paragraph("AI Insights", heading_style))
            if insights:
                for insight in insights:
                    severity_colors = {
                        "critical": colors.HexColor('#e63946'),
                        "high": colors.HexColor('#f4a261'),
                        "medium": colors.HexColor('#e9c46a'),
                        "low": colors.HexColor('#2a9d8f'),
                    }
                    sev = insight.get("severity", "medium")
                    elements.append(Paragraph(
                        f'<font color="{severity_colors.get(sev, colors.grey)}">[{sev.upper()}]</font> '
                        f'<b>{insight.get("title", "")}</b>',
                        body_style
                    ))
                    elements.append(Paragraph(insight.get("description", ""), body_style))
                    if insight.get("recommendation"):
                        elements.append(Paragraph(f'<i>Recommendation: {insight["recommendation"]}</i>', body_style))
                    elements.append(Spacer(1, 0.15*inch))
            else:
                elements.append(Paragraph("No significant insights generated.", body_style))
            elements.append(Spacer(1, 0.25*inch))

        if report_type in ("full", "bugs_only"):
            # Bugs Section
            elements.append(Paragraph("Detected Bugs", heading_style))
            if bugs:
                bug_summary = [["Severity", "Count"]]
                severity_counts = {}
                for bug in bugs:
                    sev = bug.get("severity", "medium")
                    severity_counts[sev] = severity_counts.get(sev, 0) + 1
                for sev in ["critical", "high", "medium", "low"]:
                    if sev in severity_counts:
                        bug_summary.append([sev.upper(), str(severity_counts[sev])])

                bug_table = Table(bug_summary, colWidths=[3*inch, 3*inch])
                bug_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e63946')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                elements.append(bug_table)
                elements.append(Spacer(1, 0.25*inch))

                for bug in bugs[:50]:  # Limit to 50 bugs in report
                    elements.append(Paragraph(
                        f'<b>[{bug.get("severity", "medium").upper()}]</b> {bug.get("title", "")}',
                        body_style
                    ))
                    elements.append(Paragraph(f'File: {bug.get("file_path", "")} (line {bug.get("line_number", "?")})', body_style))
                    if bug.get("explanation"):
                        elements.append(Paragraph(f'<i>{bug["explanation"]}</i>', body_style))
                    elements.append(Spacer(1, 0.1*inch))
            else:
                elements.append(Paragraph("No bugs detected. Great job!", body_style))

        doc.build(elements)
        return buffer.getvalue()

    def _generate_docx(self, analysis: dict, bugs: list, insights: list, report_type: str) -> bytes:
        """Generate DOCX report using python-docx."""
        doc = Document()

        # Title
        title = doc.add_heading('DevInsight — Code Quality Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'Report Type: {report_type.replace("_", " ").title()}')
        doc.add_paragraph('—' * 60)

        if report_type in ("full", "insights_only"):
            doc.add_heading('Executive Summary', level=1)
            table = doc.add_table(rows=8, cols=2)
            table.style = 'Light Grid Accent 1'

            metrics = [
                ("Total Files", str(analysis.get("total_files", 0))),
                ("Total Lines", str(analysis.get("total_lines", 0))),
                ("Avg Complexity", f'{analysis.get("avg_complexity", 0):.1f}'),
                ("Max Complexity", f'{analysis.get("max_complexity", 0):.1f}'),
                ("Tech Debt", f'{analysis.get("technical_debt_hours", 0):.1f} hours'),
                ("Risk Level", analysis.get("risk_level", "low").upper()),
                ("Risk Score", f'{analysis.get("risk_score", 0):.1f}/100'),
                ("Languages", str(len(analysis.get("language_breakdown", {})))),
            ]

            for i, (label, value) in enumerate(metrics):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = value

            doc.add_paragraph()
            doc.add_heading('AI Insights', level=1)
            for insight in insights:
                p = doc.add_paragraph()
                run = p.add_run(f'[{insight.get("severity", "medium").upper()}] ')
                run.bold = True
                p.add_run(insight.get("title", ""))
                doc.add_paragraph(insight.get("description", ""))
                if insight.get("recommendation"):
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run("Recommendation: ")
                    run2.italic = True
                    p2.add_run(insight["recommendation"])

        if report_type in ("full", "bugs_only"):
            doc.add_heading('Detected Bugs', level=1)
            for bug in bugs[:50]:
                p = doc.add_paragraph()
                run = p.add_run(f'[{bug.get("severity", "medium").upper()}] ')
                run.bold = True
                p.add_run(bug.get("title", ""))
                doc.add_paragraph(f'File: {bug.get("file_path", "")} (line {bug.get("line_number", "?")})')
                if bug.get("explanation"):
                    doc.add_paragraph(bug["explanation"])
                doc.add_paragraph()

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _generate_json(self, analysis: dict, bugs: list, insights: list, report_type: str) -> bytes:
        """Generate JSON report."""
        report = {
            "meta": {
                "generated_at": datetime.now().isoformat(),
                "report_type": report_type,
                "platform": "DevInsight",
            }
        }

        if report_type in ("full", "insights_only"):
            report["analysis"] = {
                "total_files": analysis.get("total_files", 0),
                "total_lines": analysis.get("total_lines", 0),
                "avg_complexity": analysis.get("avg_complexity", 0),
                "max_complexity": analysis.get("max_complexity", 0),
                "technical_debt_hours": analysis.get("technical_debt_hours", 0),
                "risk_level": analysis.get("risk_level", "low"),
                "risk_score": analysis.get("risk_score", 0),
                "language_breakdown": analysis.get("language_breakdown", {}),
                "hotspot_files": analysis.get("hotspot_files", []),
            }
            report["insights"] = insights

        if report_type in ("full", "bugs_only"):
            report["bugs"] = bugs

        return json.dumps(report, indent=2, default=str).encode('utf-8')
